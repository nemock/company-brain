"""Tests for the html/docx output formats and branding integration (v0.3.0 sub-piece 4)."""

from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from company_brain.render import render_mrd, render_one_pager
from company_brain.scaffold import scaffold


REPO_ROOT = Path(__file__).parent.parent
MEDDEV_VAULT = REPO_ROOT / "examples" / "meddev-fictional"


# ---------------------------------------------------------------------------
# HTML output
# ---------------------------------------------------------------------------


def test_render_mrd_html_contains_branded_css(tmp_path: Path) -> None:
    out = tmp_path / "mrd.html"
    render_mrd(
        MEDDEV_VAULT,
        output_path=out,
        generation_date=date(2026, 5, 24),
        output_format="html",
    )
    html = out.read_text()
    assert "<!doctype html>" in html
    # The example vault's primary color appears in the CSS root variables.
    assert "--cb-primary: #1f3a5f" in html
    # Page title pulled from doc meta.
    assert "<title>Marketing Requirements Document</title>" in html


def test_render_mrd_html_renders_tables(tmp_path: Path) -> None:
    """GFM table markdown should become <table> elements in HTML."""

    out = tmp_path / "mrd.html"
    render_mrd(
        MEDDEV_VAULT,
        output_path=out,
        generation_date=date(2026, 5, 24),
        output_format="html",
    )
    html = out.read_text()
    assert "<table>" in html
    assert "<th>" in html
    # The sources bibliography is a table.
    assert "fda-510k-summary" in html


def test_render_mrd_html_is_idempotent(tmp_path: Path) -> None:
    a = tmp_path / "a.html"
    b = tmp_path / "b.html"
    pinned = date(2026, 5, 24)
    render_mrd(MEDDEV_VAULT, output_path=a, generation_date=pinned, output_format="html")
    render_mrd(MEDDEV_VAULT, output_path=b, generation_date=pinned, output_format="html")
    assert a.read_bytes() == b.read_bytes()


def test_render_one_pager_html_works(tmp_path: Path) -> None:
    out = tmp_path / "op.html"
    render_one_pager(
        MEDDEV_VAULT,
        output_path=out,
        generation_date=date(2026, 5, 24),
        output_format="html",
    )
    html = out.read_text()
    assert "<title>One-pager</title>" in html


def test_render_one_pager_does_not_support_docx() -> None:
    with pytest.raises(ValueError, match="one-pager"):
        render_one_pager(MEDDEV_VAULT, output_format="docx", write=False)


# ---------------------------------------------------------------------------
# DOCX output
# ---------------------------------------------------------------------------


def test_render_mrd_docx_writes_valid_file(tmp_path: Path) -> None:
    out = tmp_path / "mrd.docx"
    render_mrd(
        MEDDEV_VAULT,
        output_path=out,
        generation_date=date(2026, 5, 24),
        output_format="docx",
    )
    # docx is a zip; python-docx round-trips it.
    doc = Document(out)
    # Should have many paragraphs (headings, body, footers).
    assert len(doc.paragraphs) > 30
    # Should have at least the IFU table, regulatory table, and sources table.
    assert len(doc.tables) >= 3
    # First paragraph is the H1.
    assert doc.paragraphs[0].text == "Marketing Requirements Document"


def test_render_mrd_docx_includes_table_data(tmp_path: Path) -> None:
    """The sources bibliography table should round-trip with its rows."""

    out = tmp_path / "mrd.docx"
    render_mrd(
        MEDDEV_VAULT,
        output_path=out,
        generation_date=date(2026, 5, 24),
        output_format="docx",
    )
    doc = Document(out)
    table_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "fda-510k-summary" in table_text
    assert "K231234" in table_text


def test_render_mrd_docx_is_idempotent_in_visible_content(tmp_path: Path) -> None:
    """docx zip bytes vary (zip metadata), but the visible paragraph and
    table text should be byte-identical across runs with the same date.

    This is the practical idempotency contract: re-running on an unchanged
    vault produces an unchanged-content document, even if the zip wrapper's
    metadata changes.
    """

    a = tmp_path / "a.docx"
    b = tmp_path / "b.docx"
    pinned = date(2026, 5, 24)
    render_mrd(MEDDEV_VAULT, output_path=a, generation_date=pinned, output_format="docx")
    render_mrd(MEDDEV_VAULT, output_path=b, generation_date=pinned, output_format="docx")

    def extract_text(path: Path) -> tuple[list[str], list[list[str]]]:
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs]
        tables = [
            [cell.text for row in table.rows for cell in row.cells]
            for table in doc.tables
        ]
        return paragraphs, tables

    assert extract_text(a) == extract_text(b)


# ---------------------------------------------------------------------------
# Branding integration into HTML
# ---------------------------------------------------------------------------


def test_branded_html_reflects_custom_colors(tmp_path: Path) -> None:
    """Overriding _branding/colors.yaml propagates into the HTML CSS."""

    import shutil

    vault = tmp_path / "vault"
    scaffold(vault, "default", init_git=False)
    branding_dir = vault / "_branding"
    branding_dir.mkdir(exist_ok=True)
    (branding_dir / "colors.yaml").write_text(
        "primary: \"#ff0066\"\nsecondary: \"#003366\"\n", encoding="utf-8"
    )

    out = tmp_path / "mrd.html"
    render_mrd(
        vault,
        output_path=out,
        generation_date=date(2026, 5, 24),
        output_format="html",
    )
    html = out.read_text()
    assert "--cb-primary: #ff0066" in html
    assert "--cb-secondary: #003366" in html


def test_html_wrapper_override_takes_precedence(tmp_path: Path) -> None:
    """Placing html-wrapper.html.j2 under _branding/templates/ wins."""

    vault = tmp_path / "vault"
    scaffold(vault, "default", init_git=False)
    overrides = vault / "_branding" / "templates"
    overrides.mkdir(parents=True, exist_ok=True)
    (overrides / "html-wrapper.html.j2").write_text(
        "<html><head><title>{{ title }}</title></head>"
        "<body><h1>CUSTOM SHELL</h1>{{ body_html }}</body></html>",
        encoding="utf-8",
    )
    out = tmp_path / "mrd.html"
    render_mrd(
        vault,
        output_path=out,
        generation_date=date(2026, 5, 24),
        output_format="html",
    )
    html = out.read_text()
    assert "CUSTOM SHELL" in html


# ---------------------------------------------------------------------------
# CLI surface guards
# ---------------------------------------------------------------------------


def test_render_mrd_rejects_unknown_format() -> None:
    with pytest.raises(ValueError, match="unknown output_format"):
        render_mrd(MEDDEV_VAULT, output_format="pdf", write=False)
