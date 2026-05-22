"""Tests for the read-only CLI helpers used by intake and atomize skills."""

from __future__ import annotations

from pathlib import Path

import pytest

from company_brain.intake_helpers import (
    ExtractError,
    ProfileLookupError,
    UnknownNodeTypeError,
    UnsupportedFormatError,
    describe_node,
    describe_profile,
    extract_text,
    to_json,
)
from company_brain.scaffold import scaffold


# ---------------------------------------------------------------------------
# describe-node
# ---------------------------------------------------------------------------


def test_describe_node_returns_pillar_spec() -> None:
    data = describe_node("pillar")
    assert data["name"] == "pillar"
    assert data["folder"] == "pillars"
    assert data["category"] == "epistemic"
    assert data["profile"] is None
    assert "extra_required_fields" in data
    assert isinstance(data["extra_required_fields"], list)


def test_describe_node_returns_competitor_spec_with_extra_fields() -> None:
    data = describe_node("competitor")
    field_names = {f["name"] for f in data["extra_required_fields"]}
    assert "legal_name" in field_names
    assert "canonical_url" in field_names


def test_describe_node_returns_profile_for_medical_device_types() -> None:
    data = describe_node("indication-for-use")
    assert data["profile"] == "medical-device"
    assert data["category"] == "profile-conditional"


def test_describe_node_returns_notes_when_present() -> None:
    data = describe_node("pillar")
    assert any("Non-goal" in n for n in data["notes"])


def test_describe_node_raises_for_unknown_type() -> None:
    with pytest.raises(UnknownNodeTypeError):
        describe_node("not-a-real-type")


def test_describe_node_json_round_trip() -> None:
    import json
    data = describe_node("source")
    serialized = to_json(data)
    parsed = json.loads(serialized)
    assert parsed["name"] == "source"


# ---------------------------------------------------------------------------
# describe-profile
# ---------------------------------------------------------------------------


def test_describe_profile_by_name() -> None:
    data = describe_profile(profile_name="medical-device")
    assert data["name"] == "medical-device"
    assert data["appends_controlled_document_footer"] is True
    assert "indication-for-use" in data["activated_profile_conditional_types"]


def test_describe_profile_default_has_no_activated_types() -> None:
    data = describe_profile(profile_name="default")
    assert data["activated_profile_conditional_types"] == []
    assert data["appends_controlled_document_footer"] is False


def test_describe_profile_active_node_types_include_pillar(tmp_path: Path) -> None:
    """Active node types list should include epistemic types regardless of profile."""

    data = describe_profile(profile_name="medical-device")
    names = {nt["name"] for nt in data["active_node_types"]}
    assert "pillar" in names
    assert "product" in names
    assert "indication-for-use" in names  # because medical-device


def test_describe_profile_default_active_node_types_exclude_meddev() -> None:
    data = describe_profile(profile_name="default")
    names = {nt["name"] for nt in data["active_node_types"]}
    assert "indication-for-use" not in names
    assert "pillar" in names


def test_describe_profile_from_vault_path(tmp_path: Path) -> None:
    result = scaffold(tmp_path / "v", "medical-device", init_git=False)
    data = describe_profile(vault_path=result.vault_path)
    assert data["name"] == "medical-device"


def test_describe_profile_raises_when_no_profile_md(tmp_path: Path) -> None:
    bare = tmp_path / "bare"
    bare.mkdir()
    with pytest.raises(ProfileLookupError):
        describe_profile(vault_path=bare)


def test_describe_profile_raises_for_unknown_name() -> None:
    with pytest.raises(ProfileLookupError):
        describe_profile(profile_name="not-a-real-profile")


def test_describe_profile_raises_without_path_or_name() -> None:
    with pytest.raises(ProfileLookupError):
        describe_profile()


# ---------------------------------------------------------------------------
# extract_text — docx
# ---------------------------------------------------------------------------


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Build a tiny .docx fixture with paragraphs and a table."""

    import docx

    doc = docx.Document()
    doc.add_heading("Sample PID", level=1)
    doc.add_paragraph("Purpose: validate that company-brain can ingest a Word doc.")
    doc.add_paragraph("Sponsor: Test Sponsor")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Risk"
    table.rows[0].cells[1].text = "Owner"
    table.rows[1].cells[0].text = "Test failure"
    table.rows[1].cells[1].text = "QA"

    out = tmp_path / "sample.docx"
    doc.save(str(out))
    return out


def test_extract_docx_returns_paragraph_text(sample_docx: Path) -> None:
    text = extract_text(sample_docx)
    assert "Sample PID" in text
    assert "Purpose:" in text
    assert "Sponsor: Test Sponsor" in text


def test_extract_docx_includes_table_content(sample_docx: Path) -> None:
    text = extract_text(sample_docx)
    assert "Risk" in text
    assert "Owner" in text
    assert "Test failure" in text
    assert "QA" in text
    # Table rows are joined with " | "
    assert "Risk | Owner" in text or "Risk | Owner\n" in text or "Risk | Owner " in text


# ---------------------------------------------------------------------------
# extract_text — pdf
# ---------------------------------------------------------------------------


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    """Generate a tiny PDF via Pillow → PDF (simplest cross-platform path)."""

    from PIL import Image, ImageDraw

    img = Image.new("RGB", (612, 200), color="white")
    d = ImageDraw.Draw(img)
    d.text((20, 20), "Hello from a test PDF.", fill="black")
    d.text((20, 60), "This page exists to exercise extract.", fill="black")

    out = tmp_path / "sample.pdf"
    img.save(str(out), "PDF", resolution=100.0)
    return out


def test_extract_pdf_runs_without_error(sample_pdf: Path) -> None:
    """Pillow-generated PDFs render text as images, so text extraction may be empty.

    We still want to verify the extraction code path doesn't crash and returns
    a string. Real-world PDFs with embedded text streams will produce content.
    """

    text = extract_text(sample_pdf)
    assert isinstance(text, str)


# ---------------------------------------------------------------------------
# extract_text — error paths
# ---------------------------------------------------------------------------


def test_extract_text_raises_for_missing_file(tmp_path: Path) -> None:
    with pytest.raises(ExtractError):
        extract_text(tmp_path / "does-not-exist.docx")


def test_extract_text_raises_for_unsupported_extension(tmp_path: Path) -> None:
    p = tmp_path / "thing.xml"
    p.write_text("<not supported />", encoding="utf-8")
    with pytest.raises(UnsupportedFormatError):
        extract_text(p)


def test_extract_text_raises_for_directory(tmp_path: Path) -> None:
    with pytest.raises(ExtractError):
        extract_text(tmp_path)
