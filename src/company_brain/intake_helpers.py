"""Read-only helpers consumed by the intake and atomize skills.

These commands let a Claude-driven skill stay aligned with the schema
without needing the schema package to be available in the skill prompt.
Everything here is **read-only**: the helpers describe the schema or
extract text from documents, but they never write node content. Node
content is written by the skills themselves.

Used by:

* ``cb describe-node <type>``  — JSON description of a node type.
* ``cb describe-profile``       — JSON description of the active profile
                                   in the vault at ``--path``.
* ``cb extract <file>``         — text extraction from .docx / .pdf.
"""

from __future__ import annotations

import json
from dataclasses import asdict
from pathlib import Path
from typing import Any

import yaml

from .schema import (
    NODE_TYPE_SPECS,
    PROFILES,
    get_active_node_types,
    get_node_type,
    get_profile,
)


# ---------------------------------------------------------------------------
# Errors
# ---------------------------------------------------------------------------


class IntakeHelperError(Exception):
    """Base class for helper-level errors raised by this module."""


class UnknownNodeTypeError(IntakeHelperError):
    pass


class ProfileLookupError(IntakeHelperError):
    pass


class ExtractError(IntakeHelperError):
    pass


class UnsupportedFormatError(ExtractError):
    pass


# ---------------------------------------------------------------------------
# describe-node
# ---------------------------------------------------------------------------


def describe_node(type_name: str) -> dict[str, Any]:
    """Return a JSON-serializable description of a node type spec.

    Raises :class:`UnknownNodeTypeError` if the type is not registered.
    """

    spec = get_node_type(type_name)
    if spec is None:
        valid = sorted(s.name for s in NODE_TYPE_SPECS)
        raise UnknownNodeTypeError(
            f"unknown node type '{type_name}'. Known types: {', '.join(valid)}"
        )
    return {
        "name": spec.name,
        "folder": spec.folder,
        "category": spec.category.value,
        "profile": spec.profile,
        "description": spec.description,
        "extra_required_fields": [
            {
                "name": f.name,
                "type": f.type.value,
                "description": f.description,
                "required": f.required,
            }
            for f in spec.extra_required_fields
        ],
        "notes": list(spec.notes),
    }


# ---------------------------------------------------------------------------
# describe-profile
# ---------------------------------------------------------------------------


def describe_profile(vault_path: Path | None = None, profile_name: str | None = None) -> dict[str, Any]:
    """Return a JSON-serializable description of a profile.

    If ``profile_name`` is provided, describe that profile directly. Otherwise
    read the active profile from ``<vault_path>/_system/PROFILE.md``.

    The returned dict includes the full list of currently-active node types
    (their names, folders, and category) so the calling skill can route
    incoming nodes to the right folder without re-reading the schema.

    Raises :class:`ProfileLookupError` on missing PROFILE.md or unknown profile.
    """

    if profile_name is None:
        if vault_path is None:
            raise ProfileLookupError(
                "must provide either vault_path or profile_name"
            )
        profile_md = vault_path / "_system" / "PROFILE.md"
        if not profile_md.is_file():
            raise ProfileLookupError(
                f"no _system/PROFILE.md in {vault_path}"
            )
        text = profile_md.read_text(encoding="utf-8")
        fm = _read_yaml_frontmatter(text)
        profile_name = fm.get("profile") if fm else None
        if not profile_name:
            raise ProfileLookupError(
                f"_system/PROFILE.md in {vault_path} has no `profile` field"
            )

    profile = get_profile(profile_name)
    if profile is None:
        known = sorted(PROFILES.keys())
        raise ProfileLookupError(
            f"unknown profile '{profile_name}'. Known: {', '.join(known)}"
        )

    active = get_active_node_types(profile_name)
    return {
        "name": profile.name,
        "description": profile.description,
        "appends_controlled_document_footer": profile.appends_controlled_document_footer,
        "notes": list(profile.notes),
        "activated_profile_conditional_types": list(profile.activated_node_type_names),
        "active_node_types": [
            {
                "name": spec.name,
                "folder": spec.folder,
                "category": spec.category.value,
                "profile": spec.profile,
            }
            for spec in active
        ],
    }


def _read_yaml_frontmatter(text: str) -> dict[str, Any] | None:
    if not text.startswith("---"):
        return None
    rest = text[3:]
    end = rest.find("\n---")
    if end == -1:
        return None
    fm_text = rest[:end].strip("\n")
    try:
        data = yaml.safe_load(fm_text)
    except yaml.YAMLError:
        return None
    return data if isinstance(data, dict) else None


# ---------------------------------------------------------------------------
# extract
# ---------------------------------------------------------------------------


def extract_text(file_path: Path) -> str:
    """Extract text content from a binary document.

    Supports:
      * .docx  (via python-docx)
      * .pdf   (via pdfplumber; text + simple tables)

    Plain text formats (.md, .txt, transcripts) are not handled here — the
    calling skill can read those directly. Image formats (.png, .jpg) are
    handled by Claude's native vision; not extracted here.

    Raises :class:`ExtractError` on any failure;
    :class:`UnsupportedFormatError` for unrecognized extensions.
    """

    if not file_path.exists():
        raise ExtractError(f"{file_path} does not exist")
    if not file_path.is_file():
        raise ExtractError(f"{file_path} is not a file")

    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return _extract_docx(file_path)
    if suffix == ".pdf":
        return _extract_pdf(file_path)
    raise UnsupportedFormatError(
        f"unsupported extension '{suffix}'. Supported: .docx, .pdf. "
        f"For markdown / plain text / transcripts, read the file directly. "
        f"For images, use Claude's vision capabilities."
    )


def _extract_docx(file_path: Path) -> str:
    try:
        import docx
    except ImportError as exc:
        raise ExtractError(
            "python-docx is required for .docx extraction. "
            "Reinstall company-brain: `uv tool install --reinstall .`"
        ) from exc

    document = docx.Document(str(file_path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            if any(row_cells):
                parts.append(" | ".join(row_cells))

    return "\n\n".join(parts) + ("\n" if parts else "")


def _extract_pdf(file_path: Path) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise ExtractError(
            "pdfplumber is required for .pdf extraction. "
            "Reinstall company-brain: `uv tool install --reinstall .`"
        ) from exc

    parts: list[str] = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            page_text = page_text.strip()
            if page_text:
                parts.append(f"--- page {page_num} ---")
                parts.append(page_text)

            # Append simple table content if present.
            tables = page.extract_tables() or []
            for tbl_idx, tbl in enumerate(tables, start=1):
                parts.append(f"--- page {page_num} table {tbl_idx} ---")
                for row in tbl:
                    cleaned = [(cell or "").strip() for cell in row]
                    if any(cleaned):
                        parts.append(" | ".join(cleaned))

    return "\n\n".join(parts) + ("\n" if parts else "")


def to_json(obj: dict[str, Any], indent: int = 2) -> str:
    """Render a dict as deterministically-keyed JSON for CLI consumption."""

    # asdict-like processing is not needed because describe_* return plain dicts.
    _ = asdict  # keep the import wired; suppresses unused-import noise
    return json.dumps(obj, indent=indent, sort_keys=False, default=str)
