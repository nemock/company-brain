"""Branded DOCX output for generated documents.

A small markdown-to-docx converter that covers the shapes the v0.3.0
generators actually emit:

* Headings (``#``, ``##``, ``###``)
* Paragraphs (plain text + inline ``**bold**`` and ``` `code` ```)
* Blockquotes (lines starting with ``> ``)
* Unordered lists (lines starting with ``- ``)
* GFM-style tables (``|`` separated rows + a ``|---|`` divider)
* Horizontal rule (``---`` on its own line)

Notably **does not** support: ordered lists, nested lists, fenced code
blocks, images, links (links pass through as plain text). The MRD and
one-pager templates were designed not to need those. If a future
generator does, extend here rather than reaching for pandoc — we want
``uv tool install`` to be enough to use ``cb render``.

Branding integration: H1/H2/H3 colors come from the branding primary
color; body font name from ``font_family_body`` (the font file itself is
not embedded — docx renders with whatever font the reader has installed).
"""

from __future__ import annotations

import re
import zipfile
from datetime import date, datetime, timezone
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .engine import Branding


# Earliest legal zip timestamp. Used to normalize every entry in the
# rendered .docx so the bytes are deterministic across runs.
_EPOCH_ZIP_TIMESTAMP = (1980, 1, 1, 0, 0, 0)
# Stable values for the docx core properties that python-docx would
# otherwise fill with the current time / the OS user — both sources of
# non-determinism in the visible Word "Modified" date and the
# `<dc:creator>` element.
_DETERMINISTIC_AUTHOR = "company-brain"


_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")
_CODE_PATTERN = re.compile(r"`([^`]+?)`")


def render_docx(
    *,
    title: str,
    body_markdown: str,
    branding: Branding,
    output_path: Path | None = None,
    generation_date: date | None = None,
) -> bytes:
    """Convert ``body_markdown`` to a branded .docx document.

    Returns the document bytes; also writes to ``output_path`` when given.

    Output is **byte-deterministic** when ``generation_date`` is pinned.
    The docx core properties ``<dcterms:created>`` and
    ``<dcterms:modified>`` are stamped to ``generation_date`` at midnight
    UTC; the inner zip is then rewritten with normalized entry order
    and a fixed (1980-01-01) timestamp per entry. Re-running on the
    same inputs produces an identical byte string — matching the
    idempotency contract for the markdown / html outputs.
    """

    doc = Document()
    _apply_default_styles(doc, branding)
    _apply_deterministic_core_properties(
        doc, generation_date or date.today()
    )

    primary_rgb = _hex_to_rgb(branding.primary)
    secondary_rgb = _hex_to_rgb(branding.secondary)

    lines = body_markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            _add_heading(doc, stripped[2:].strip(), level=1, color=primary_rgb)
        elif stripped.startswith("## "):
            _add_heading(doc, stripped[3:].strip(), level=2, color=primary_rgb)
        elif stripped.startswith("### "):
            _add_heading(doc, stripped[4:].strip(), level=3, color=secondary_rgb)
        elif stripped.startswith("> "):
            # Consume a run of blockquote lines into one block.
            block_lines: list[str] = []
            while i < len(lines) and lines[i].rstrip().startswith("> "):
                block_lines.append(lines[i].rstrip()[2:])
                i += 1
            _add_blockquote(doc, " ".join(block_lines))
            continue
        elif stripped.startswith("- "):
            # Consume a run of bullet items.
            while i < len(lines) and lines[i].rstrip().startswith("- "):
                _add_bullet(doc, lines[i].rstrip()[2:])
                i += 1
            continue
        elif stripped == "---":
            _add_horizontal_rule(doc)
        elif _looks_like_table_row(stripped):
            # Collect contiguous table rows.
            table_rows: list[str] = []
            while i < len(lines) and _looks_like_table_row(lines[i].rstrip()):
                table_rows.append(lines[i].rstrip())
                i += 1
            _add_table(doc, table_rows, primary_rgb)
            continue
        else:
            _add_paragraph(doc, stripped)
        i += 1

    buffer = BytesIO()
    doc.save(buffer)
    data = _normalize_docx_zip(buffer.getvalue())
    if output_path is not None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_bytes(data)
    return data


# ---------------------------------------------------------------------------
# Block emitters
# ---------------------------------------------------------------------------


def _add_heading(doc, text: str, *, level: int, color: RGBColor) -> None:
    para = doc.add_heading(level=level)
    run = para.add_run(text)
    run.font.color.rgb = color


def _add_paragraph(doc, text: str) -> None:
    para = doc.add_paragraph()
    _add_styled_runs(para, text)


def _add_blockquote(doc, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Pt(18)
    para.paragraph_format.right_indent = Pt(18)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.italic = True


def _add_bullet(doc, text: str) -> None:
    para = doc.add_paragraph(style="List Bullet")
    _add_styled_runs(para, text)


def _add_horizontal_rule(doc) -> None:
    para = doc.add_paragraph()
    run = para.add_run("─" * 40)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_table(doc, rows: list[str], header_rgb: RGBColor) -> None:
    parsed_rows = [_parse_table_row(r) for r in rows]
    # A GFM table is: header row, divider row, then body rows. Filter the
    # divider out (cells that look like ---).
    body_rows = [r for r in parsed_rows if not _is_divider_row(r)]
    if not body_rows:
        return
    headers = body_rows[0]
    data_rows = body_rows[1:]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        para = header_cells[idx].paragraphs[0]
        run = para.add_run(header)
        run.bold = True
        run.font.color.rgb = header_rgb
    for data_row in data_rows:
        # Pad if a row is short.
        row_cells = table.add_row().cells
        for idx in range(len(headers)):
            text = data_row[idx] if idx < len(data_row) else ""
            cell_para = row_cells[idx].paragraphs[0]
            _add_styled_runs(cell_para, text)


def _add_styled_runs(para, text: str) -> None:
    """Render text with **bold** and `code` inline spans into ``para``."""

    # Walk the text, splitting on bold-or-code spans. We do this in a single
    # pass by finding the earliest match each step.
    remaining = text
    while remaining:
        bold_match = _BOLD_PATTERN.search(remaining)
        code_match = _CODE_PATTERN.search(remaining)
        next_match = _earliest(bold_match, code_match)
        if next_match is None:
            para.add_run(remaining)
            return
        if next_match.start() > 0:
            para.add_run(remaining[: next_match.start()])
        run = para.add_run(next_match.group(1))
        if next_match is bold_match:
            run.bold = True
        else:
            run.font.name = "Courier New"
        remaining = remaining[next_match.end():]


def _earliest(a, b):
    if a is None:
        return b
    if b is None:
        return a
    return a if a.start() <= b.start() else b


# ---------------------------------------------------------------------------
# Table parsing
# ---------------------------------------------------------------------------


def _looks_like_table_row(text: str) -> bool:
    """A line starting and ending with ``|`` (after trimming) is a GFM row."""

    s = text.strip()
    return s.startswith("|") and s.endswith("|") and s.count("|") >= 2


def _parse_table_row(text: str) -> list[str]:
    """Split a ``| a | b | c |`` row into cells. Trims each cell."""

    inner = text.strip().strip("|")
    return [c.strip() for c in inner.split("|")]


_DIVIDER_RE = re.compile(r"^:?-{2,}:?$")


def _is_divider_row(cells: list[str]) -> bool:
    return all(_DIVIDER_RE.match(c) for c in cells) and len(cells) > 0


# ---------------------------------------------------------------------------
# Style setup
# ---------------------------------------------------------------------------


def _apply_default_styles(doc, branding: Branding) -> None:
    """Set the document-default font names from branding.

    The font *file* is not embedded (would require shipping a font asset);
    Word renders with whatever font the reader has installed. Most clients
    will substitute gracefully when the named font is missing.
    """

    body_font = _strip_family(branding.font_family_body)
    style = doc.styles["Normal"]
    style.font.name = body_font
    style.font.size = Pt(11)


def _strip_family(font_family: str) -> str:
    """Pull the first font name out of a CSS-style family string."""

    first = font_family.split(",")[0].strip()
    return first.strip('"').strip("'")


# ---------------------------------------------------------------------------
# Misc
# ---------------------------------------------------------------------------


def _apply_deterministic_core_properties(doc, generation_date: date) -> None:
    """Stamp the docx core properties with deterministic values.

    Without this, python-docx defaults ``<dcterms:created>`` and
    ``<dcterms:modified>`` to the current wall-clock time on the
    machine running the render. That makes the bytes vary on every run
    even when the body content is identical.

    Map ``generation_date`` to midnight UTC so the visible "Modified"
    date in Word lines up with the footer text the markdown writer
    already emits, and so the byte sequence inside ``docProps/core.xml``
    is fully determined by the inputs.
    """

    pinned = datetime(
        generation_date.year,
        generation_date.month,
        generation_date.day,
        0,
        0,
        0,
        tzinfo=timezone.utc,
    )
    cp = doc.core_properties
    cp.author = _DETERMINISTIC_AUTHOR
    cp.last_modified_by = _DETERMINISTIC_AUTHOR
    cp.created = pinned
    cp.modified = pinned


def _normalize_docx_zip(raw: bytes) -> bytes:
    """Rewrite a docx zip with deterministic timestamps and entry order.

    A .docx is a zip. Python's zipfile module stamps every entry with
    the current local time and external_attr derived from the running
    OS — both vary per run. Re-pack each entry with the (1980-01-01)
    epoch timestamp, zero external_attr, and a stable name-sorted
    order. The visible content of each entry is preserved byte-for-byte.
    """

    in_buf = BytesIO(raw)
    out_buf = BytesIO()
    with zipfile.ZipFile(in_buf, "r") as zin:
        entries = sorted(zin.infolist(), key=lambda info: info.filename)
        with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in entries:
                data = zin.read(info.filename)
                new_info = zipfile.ZipInfo(filename=info.filename)
                new_info.date_time = _EPOCH_ZIP_TIMESTAMP
                # Reuse the compress type from the source so we don't
                # accidentally re-compress a stored entry. external_attr
                # captures unix file mode + DOS attrs; zero it for
                # cross-platform stability.
                new_info.compress_type = info.compress_type
                new_info.external_attr = 0
                zout.writestr(new_info, data)
    return out_buf.getvalue()


def _hex_to_rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    if len(h) != 6:
        return RGBColor(0x1F, 0x3A, 0x5F)  # fall back to default primary
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
